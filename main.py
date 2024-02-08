import os
import json
import copy
import random
import argparse
import datetime
from typing import (
    Dict,
    Union
)

import docx
from docx.text.run import Run
from docx.enum.text import WD_ALIGN_PARAGRAPH


def get_month_year_map(start_year, end_year):
    return {
        "OTTOBRE": start_year,
        "NOVEMBRE": start_year,
        "DICEMBRE": start_year,
        "GENNAIO": end_year,
        "FEBBRAIO": end_year,
        "MARZO": end_year,
        "APRILE": end_year,
        "MAGGIO": end_year,
        "GIUGNO": end_year,
        "LUGLIO": end_year,
        "AGOSTO": end_year,
        "SETTEMBRE": end_year
    }


def get_month_idx(month):
    return [
        "GENNAIO",
        "FEBBRAIO",
        "MARZO",
        "APRILE",
        "MAGGIO",
        "GIUGNO",
        "LUGLIO",
        "AGOSTO",
        "SETTEMBRE",
        "OTTOBRE",
        "NOVEMBRE",
        "DICEMBRE"
    ].index(month.upper()) + 1


def add_paragraph_text_with_style(paragraph, new_text, base_style):
    style = copy.deepcopy(base_style)
    child = paragraph._p._insert_r(style)
    r = Run(child, paragraph)
    r.text = new_text
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def parse_month_data(month_data, dates_sep="-"):
    new_month_data = {}
    for key, value in month_data.items():
        if key.lower() == "default":
            new_month_data[key] = value if isinstance(value[0], list) else [value]
        elif dates_sep in key:
            start_date, end_date = map(int, key.split(dates_sep))
            for i in range(start_date, end_date + 1):
                new_month_data[str(i)] = value
        else:
            new_month_data[key] = value

    return new_month_data


def modify_docx(document_path: Union[str, os.PathLike],
                month_data: Dict,
                current_year: str,
                current_month: str,
                default: str = None,
                hour_format: str = "%H:%M"):
    default = default or [["ANALISI LETTERATURA", "10:00", "16:00"]]

    month_data = parse_month_data(month_data)
    document = docx.Document(document_path)

    month_year_table = document.tables[0]
    activity_table = document.tables[1]

    base_style = month_year_table.rows[0].cells[0].paragraphs[0].runs[0]._r

    add_paragraph_text_with_style(month_year_table.rows[1].cells[0].paragraphs[0], current_year, base_style)
    add_paragraph_text_with_style(month_year_table.rows[1].cells[1].paragraphs[0], current_month, base_style)

    for row in activity_table.rows[1:]:
        day = row.cells[0].text
        try:
            date = datetime.datetime(year=int(current_year), month=get_month_idx(current_month), day=int(day))
            weekday = date.weekday()
        except ValueError:
            # The date is not valid, e.g., February 31
            break

        if weekday == 6:
            continue

        if day in month_data:
            activity, start_hour, end_hour = month_data[day]
        else:
            curr_default = month_data["default"] if "default" in month_data else default
            activity, start_hour, end_hour = random.choice(curr_default)

        add_paragraph_text_with_style(row.cells[1].paragraphs[0], activity, base_style)
        add_paragraph_text_with_style(row.cells[2].paragraphs[0], start_hour, base_style)
        add_paragraph_text_with_style(row.cells[3].paragraphs[0], end_hour, base_style)

        start_dt = datetime.datetime.strptime(start_hour, hour_format)
        end_dt = datetime.datetime.strptime(end_hour, hour_format)
        duration = abs(start_dt - end_dt)
        tot_hours = str(duration)[:-3]
        tot_hours = tot_hours if not tot_hours.endswith('00') else tot_hours.split(':')[0]
        add_paragraph_text_with_style(row.cells[4].paragraphs[0], tot_hours, base_style)

    return document


def generate_registries(document_path: Union[str, os.PathLike],
                        registry_data: Dict,
                        year_sep='-'):
    for year_interval, year_data in registry_data.items():
        start_year, end_year = year_interval.split(year_sep)
        year_map = get_month_year_map(start_year, end_year)

        year_document = None
        n_months = len(year_data)
        for i, (month, month_data) in enumerate(year_data.items()):
            month_document = modify_docx(
                document_path,
                month_data,
                year_map[month],
                month
            )

            if i < n_months - 1:
                month_document.add_page_break()

            if i == 0:
                year_document = month_document
            else:
                for element in month_document.element.body:
                    year_document.element.body.append(element)

        yield year_interval, year_document


if __name__ == "__main__":
    script_path = os.path.dirname(os.path.realpath(__file__))

    parser = argparse.ArgumentParser()
    parser.add_argument('document', type=str)
    parser.add_argument('data', type=str)
    parser.add_argument('--registry_path', '-rp', type=str, default=None)
    parser.add_argument('--year_path_suffix', '-yps', type=str, default="registro_dottorato")
    parser.add_argument('--year_separator', '-ys', type=str, default="-")
    args = parser.parse_args()

    if args.registry_path is None:
        args.registry_path = os.path.join(script_path, "registry_per_year")
    os.makedirs(args.registry_path, exist_ok=True)

    if not all(os.path.exists(os.path.join(script_path, p)) for p in [args.document, args.data, args.registry_path]):
        raise FileNotFoundError("One of the input files does not exist")

    with open(args.data, 'r') as json_data:
        data = json.load(json_data)

    generator = generate_registries(
        args.document,
        data,
        year_sep=args.year_separator
    )

    for year, year_registry in generator:
        year_registry.save(
            os.path.join(args.registry_path, f"{year}_{args.year_path_suffix}.docx")
        )
